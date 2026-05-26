"""Pre-process quality check for documents bound for Azure AI Content Understanding.

This is the in-process Python API used by `function_app.py` before a raw
document is submitted to Content Understanding. Same checks and same severity
model as the original standalone CLI prototype — minus the CLI, argparse and
SQLite persistence (we persist to Azure SQL via `sql_client.py`).

Test fixtures (sample PDFs covering pass / warn / fail cases) and the
generator script live in `.e2e/samples/` and `.e2e/generate_sample_pdfs.py`.

Public surface:
    Severity          — ERROR | WARNING | INFO
    QualityIssue      — one finding (code + severity + message + details)
    QualityReport     — full result for one file (passed, score, band, issues, metadata)
    check_document()  — run all applicable checks against a file path

The checker is conservative: it flags anything that would be rejected outright
(ERROR) plus anything likely to produce poor extraction quality (WARNING).
Optional dependencies (pymupdf, pillow, python-docx, openpyxl, python-pptx)
are loaded lazily — if missing, the corresponding deep checks are skipped and
an INFO issue is recorded.
"""

from __future__ import annotations

import mimetypes
import os
from dataclasses import asdict, dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any, Optional


# ---------------------------------------------------------------------------
# Limits (kept as module-level constants so they're easy to audit / override).
# ---------------------------------------------------------------------------

MB = 1024 * 1024

# Standard mode limits
DOC_IMAGE_MAX_BYTES = 200 * MB
DOC_IMAGE_MAX_PAGES = 300
OFFICE_MAX_BYTES = 200 * MB
OFFICE_MAX_CHARS = 1_000_000
TEXT_MAX_BYTES = 1 * MB
TEXT_MAX_CHARS = 1_000_000
IMAGE_MIN_DIM = 50
IMAGE_MAX_DIM = 10_000

# Pro mode (preview) limits
PRO_MAX_BYTES = 100 * MB
PRO_MAX_PAGES = 150
PRO_SUPPORTED_EXT = {".pdf", ".tiff", ".tif", ".jpg", ".jpeg", ".jpe",
                     ".png", ".bmp", ".heif", ".heic"}

# Extension categories
EXT_PDF = {".pdf"}
EXT_IMAGE = {".jpg", ".jpeg", ".jpe", ".png", ".bmp", ".heif", ".heic",
             ".tiff", ".tif"}
EXT_OFFICE = {".docx", ".xlsx", ".pptx"}
EXT_TEXT = {".txt", ".html", ".htm", ".md", ".rtf", ".eml", ".msg", ".xml"}

ALL_SUPPORTED_EXT = EXT_PDF | EXT_IMAGE | EXT_OFFICE | EXT_TEXT

# Page-equivalent rules (used for billing visibility, not hard limits)
CHARS_PER_TEXT_PAGE = 3_000

# ---- Document quality thresholds (tune to taste) -----------------------------
LOW_TEXT_CHAR_THRESHOLD = 50
MIN_AVG_CHARS_PER_PAGE = 200
LOW_TEXT_PAGE_RATIO = 0.30
PDF_PAGE_SCAN_CAP = 100
IMAGE_BLUR_VARIANCE_MIN = 50.0
IMAGE_LOW_CONTRAST_STDEV = 15.0

# ---- Scanned-PDF quality thresholds ------------------------------------------
PDF_SCAN_BLUR_VARIANCE_MIN = 200.0
PDF_SCAN_LOW_CONTRAST_STDEV = 12.0
PDF_SCAN_PAGES_TO_RENDER = 5
PDF_SCAN_RENDER_DPI = 150
PDF_SCAN_SKEW_THRESHOLD_DEG = 2.0
PDF_SCAN_SKEW_MAX_DEG = 10
PDF_SCAN_SKEW_CONFIDENCE_MIN = 1.15
PDF_SCAN_MIN_DPI = 150.0

# Magic-byte signatures (extension/content sanity check).
MAGIC_SIGNATURES: dict[str, tuple[bytes, ...]] = {
    ".pdf":  (b"%PDF-",),
    ".png":  (b"\x89PNG\r\n\x1a\n",),
    ".jpg":  (b"\xff\xd8\xff",),
    ".jpeg": (b"\xff\xd8\xff",),
    ".jpe":  (b"\xff\xd8\xff",),
    ".bmp":  (b"BM",),
    ".tif":  (b"II*\x00", b"MM\x00*"),
    ".tiff": (b"II*\x00", b"MM\x00*"),
    ".heif": (b"ftypheic", b"ftypheix", b"ftyphevc", b"ftypmif1", b"ftypmsf1"),
    ".heic": (b"ftypheic", b"ftypheix", b"ftyphevc", b"ftypmif1", b"ftypmsf1"),
    # ZIP-based Office formats
    ".docx": (b"PK\x03\x04",),
    ".xlsx": (b"PK\x03\x04",),
    ".pptx": (b"PK\x03\x04",),
}


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

class Severity(str, Enum):
    ERROR = "ERROR"
    WARNING = "WARNING"
    INFO = "INFO"


@dataclass
class QualityIssue:
    code: str
    severity: Severity
    message: str
    details: dict[str, Any] = field(default_factory=dict)


# Scoring: start at 100 and subtract per-issue penalties, clamp to [0, 100].
SCORE_PENALTY = {
    Severity.ERROR: 30,
    Severity.WARNING: 10,
    Severity.INFO: 2,
}

SCORE_BANDS: list[tuple[int, str]] = [
    (90, "excellent"),
    (75, "good"),
    (50, "fair"),
    (25, "poor"),
    (0, "unusable"),
]


def _compute_score(issues: list["QualityIssue"]) -> int:
    score = 100
    for issue in issues:
        score -= SCORE_PENALTY.get(issue.severity, 0)
    return max(0, min(100, score))


def _score_band(score: int) -> str:
    for threshold, label in SCORE_BANDS:
        if score >= threshold:
            return label
    return "unusable"


@dataclass
class QualityReport:
    file_path: str
    file_size_bytes: int
    extension: str
    detected_kind: str        # "pdf" | "image" | "office" | "text" | "unknown"
    mode: str                 # "standard" | "pro"
    passed: bool
    issues: list[QualityIssue] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)

    def add(self, issue: QualityIssue) -> None:
        self.issues.append(issue)
        if issue.severity is Severity.ERROR:
            self.passed = False

    @property
    def score(self) -> int:
        return _compute_score(self.issues)

    @property
    def band(self) -> str:
        return _score_band(self.score)

    @property
    def error_count(self) -> int:
        return sum(1 for i in self.issues if i.severity is Severity.ERROR)

    @property
    def warning_count(self) -> int:
        return sum(1 for i in self.issues if i.severity is Severity.WARNING)

    @property
    def info_count(self) -> int:
        return sum(1 for i in self.issues if i.severity is Severity.INFO)

    def to_dict(self) -> dict[str, Any]:
        d = asdict(self)
        for i in d["issues"]:
            if isinstance(i["severity"], Severity):
                i["severity"] = i["severity"].value
        d["score"] = self.score
        d["band"] = self.band
        d["error_count"] = self.error_count
        d["warning_count"] = self.warning_count
        d["info_count"] = self.info_count
        return d


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _detect_kind(ext: str) -> str:
    if ext in EXT_PDF:
        return "pdf"
    if ext in EXT_IMAGE:
        return "image"
    if ext in EXT_OFFICE:
        return "office"
    if ext in EXT_TEXT:
        return "text"
    return "unknown"


def _read_head(path: Path, n: int = 16) -> bytes:
    with path.open("rb") as f:
        return f.read(n)


def _check_magic(path: Path, ext: str, report: QualityReport) -> None:
    sigs = MAGIC_SIGNATURES.get(ext)
    if not sigs:
        return
    try:
        head = _read_head(path, 32)
    except OSError as exc:
        report.add(QualityIssue(
            code="READ_FAILED",
            severity=Severity.ERROR,
            message=f"Could not read file: {exc}",
        ))
        return

    if ext in {".heif", ".heic"}:
        ok = any(sig in head for sig in sigs)
    else:
        ok = any(head.startswith(sig) for sig in sigs)

    if not ok:
        report.add(QualityIssue(
            code="MAGIC_MISMATCH",
            severity=Severity.WARNING,
            message=(f"File header does not match the {ext} extension. "
                     "The file may be renamed or corrupted."),
            details={"first_bytes": head[:8].hex()},
        ))


# ---------------------------------------------------------------------------
# Per-kind checks
# ---------------------------------------------------------------------------

def _estimate_page_skew(gray_img, max_angle: int = PDF_SCAN_SKEW_MAX_DEG
                        ) -> tuple[int, float]:
    """Estimate text-line skew (degrees) via projection-profile variance."""
    from PIL import Image

    w, h = gray_img.size
    target_w = 200
    if w > target_w:
        ratio = target_w / w
        small = gray_img.resize((target_w, max(1, int(h * ratio))),
                                resample=Image.BILINEAR)
    else:
        small = gray_img
    binary = small.point(lambda p: 0 if p < 128 else 255)

    best_var = -1.0
    best_angle = 0
    var_at_zero = 0.0
    for angle in range(-max_angle, max_angle + 1):
        rot = binary.rotate(angle, resample=Image.NEAREST, fillcolor=255)
        data = rot.tobytes()
        rw, rh = rot.size
        rows = [255 * rw - sum(data[y * rw:(y + 1) * rw])
                for y in range(rh)]
        mean = sum(rows) / rh
        var = sum((r - mean) ** 2 for r in rows) / rh
        if angle == 0:
            var_at_zero = var
        if var > best_var:
            best_var = var
            best_angle = angle
    confidence = best_var / max(var_at_zero, 1.0)
    return best_angle, confidence


def _estimate_page_image_dpi(page) -> float | None:
    """Return the minimum effective DPI of the largest image on a PDF page."""
    best_area = 0.0
    best_dpi: float | None = None
    try:
        infos = page.get_image_info()
    except Exception:  # noqa: BLE001 - older PyMuPDF or odd PDFs
        return None
    for info in infos:
        bbox = info.get("bbox")
        w_px = info.get("width", 0)
        h_px = info.get("height", 0)
        if not bbox or not w_px or not h_px:
            continue
        x0, y0, x1, y1 = bbox
        bw_pt = abs(x1 - x0)
        bh_pt = abs(y1 - y0)
        if bw_pt <= 0 or bh_pt <= 0:
            continue
        area = bw_pt * bh_pt
        if area > best_area:
            best_area = area
            dpi_x = w_px / (bw_pt / 72.0)
            dpi_y = h_px / (bh_pt / 72.0)
            best_dpi = min(dpi_x, dpi_y)
    return best_dpi


def _analyze_pdf_scan_quality(doc, low_text_indices: list[int],
                              report: QualityReport) -> None:
    """Render image-likely PDF pages and check for blur / low contrast / skew."""
    if not low_text_indices:
        return
    try:
        import fitz  # PyMuPDF
        from PIL import Image, ImageFilter, ImageStat
    except ImportError:
        report.add(QualityIssue(
            code="PDF_SCAN_CHECK_SKIPPED",
            severity=Severity.INFO,
            message=("Pillow is not installed; cannot check scanned-page "
                     "image quality. Install with: pip install pillow"),
        ))
        return

    n = min(len(low_text_indices), PDF_SCAN_PAGES_TO_RENDER)
    step = len(low_text_indices) / n
    sample_indices = [low_text_indices[int(i * step)] for i in range(n)]

    blur_vars: list[float] = []
    stdevs: list[float] = []
    means: list[float] = []
    blurry_pages: list[int] = []
    low_contrast_pages: list[int] = []
    skewed_pages: list[tuple[int, int]] = []
    low_dpi_pages: list[tuple[int, float]] = []
    dpis: list[float] = []

    laplacian = ImageFilter.Kernel(
        size=(3, 3),
        kernel=(0, 1, 0, 1, -4, 1, 0, 1, 0),
        scale=1,
        offset=0,
    )

    for idx in sample_indices:
        try:
            page = doc.load_page(idx)
            pix = page.get_pixmap(dpi=PDF_SCAN_RENDER_DPI,
                                  colorspace=fitz.csGRAY)
            img = Image.frombytes("L", (pix.width, pix.height), pix.samples)
            edges = img.filter(laplacian)
            blur_var = ImageStat.Stat(edges).var[0]
            stat = ImageStat.Stat(img)
            stdev = stat.stddev[0]
            mean = stat.mean[0]
            if stdev >= PDF_SCAN_LOW_CONTRAST_STDEV:
                skew, skew_confidence = _estimate_page_skew(img)
            else:
                skew, skew_confidence = 0, 1.0
            page_dpi = _estimate_page_image_dpi(page)
        except Exception:  # noqa: BLE001 - one bad page should not abort
            continue
        blur_vars.append(blur_var)
        stdevs.append(stdev)
        means.append(mean)
        if stdev < PDF_SCAN_LOW_CONTRAST_STDEV:
            low_contrast_pages.append(idx + 1)
        elif blur_var < PDF_SCAN_BLUR_VARIANCE_MIN:
            blurry_pages.append(idx + 1)
        if (abs(skew) >= PDF_SCAN_SKEW_THRESHOLD_DEG
                and skew_confidence >= PDF_SCAN_SKEW_CONFIDENCE_MIN):
            skewed_pages.append((idx + 1, skew))
        if page_dpi is not None:
            dpis.append(page_dpi)
            if page_dpi < PDF_SCAN_MIN_DPI:
                low_dpi_pages.append((idx + 1, page_dpi))

    if not blur_vars:
        return

    avg_blur = sum(blur_vars) / len(blur_vars)
    avg_stdev = sum(stdevs) / len(stdevs)
    avg_mean = sum(means) / len(means)

    report.metadata["scan_pages_analyzed"] = len(blur_vars)
    report.metadata["scan_avg_blur_variance"] = round(avg_blur, 1)
    report.metadata["scan_avg_stdev"] = round(avg_stdev, 1)
    report.metadata["scan_avg_mean_intensity"] = round(avg_mean, 1)

    if blurry_pages:
        report.add(QualityIssue(
            code="PDF_SCAN_BLURRY",
            severity=Severity.ERROR,
            message=(f"{len(blurry_pages)} of {len(blur_vars)} rendered "
                     f"scanned pages look blurry or smudged "
                     f"(avg variance of Laplacian = {avg_blur:.1f}, "
                     f"threshold {PDF_SCAN_BLUR_VARIANCE_MIN}). "
                     "OCR accuracy will be poor; rescan at higher quality."),
            details={"blurry_pages": blurry_pages,
                     "pages_analyzed": len(blur_vars),
                     "avg_blur_variance": round(avg_blur, 1),
                     "threshold": PDF_SCAN_BLUR_VARIANCE_MIN},
        ))

    if low_contrast_pages:
        if avg_mean > 230:
            why = "pages are nearly white (faded or under-exposed scan)"
        elif avg_mean < 25:
            why = "pages are nearly black (over-exposed scan)"
        else:
            why = "pixel intensities have very little variation"
        report.add(QualityIssue(
            code="PDF_SCAN_LOW_CONTRAST",
            severity=Severity.ERROR,
            message=(f"{len(low_contrast_pages)} of {len(blur_vars)} rendered "
                     f"scanned pages have low contrast: {why} "
                     f"(avg stdev = {avg_stdev:.1f}, threshold "
                     f"{PDF_SCAN_LOW_CONTRAST_STDEV}). Text extraction may fail."),
            details={"low_contrast_pages": low_contrast_pages,
                     "pages_analyzed": len(blur_vars),
                     "avg_stdev": round(avg_stdev, 1),
                     "avg_mean_intensity": round(avg_mean, 1)},
        ))

    if skewed_pages:
        max_skew = max(abs(a) for _, a in skewed_pages)
        report.metadata["scan_max_skew_deg"] = max_skew
        report.add(QualityIssue(
            code="PDF_SCAN_SKEWED",
            severity=Severity.ERROR,
            message=(f"{len(skewed_pages)} of {len(blur_vars)} rendered "
                     f"scanned pages are tilted by >={PDF_SCAN_SKEW_THRESHOLD_DEG} deg "
                     f"(max {max_skew} deg). Skewed scans cause OCR to misread "
                     "or split lines; deskew the source before submitting."),
            details={"skewed_pages": [{"page": p, "angle_deg": a}
                                      for p, a in skewed_pages],
                     "pages_analyzed": len(blur_vars),
                     "threshold_deg": PDF_SCAN_SKEW_THRESHOLD_DEG},
        ))

    if low_dpi_pages:
        min_dpi = min(d for _, d in low_dpi_pages)
        avg_dpi = sum(dpis) / len(dpis) if dpis else 0.0
        report.metadata["scan_min_image_dpi"] = round(min_dpi, 1)
        report.metadata["scan_avg_image_dpi"] = round(avg_dpi, 1)
        report.add(QualityIssue(
            code="PDF_SCAN_LOW_DPI",
            severity=Severity.ERROR,
            message=(f"{len(low_dpi_pages)} of {len(blur_vars)} rendered "
                     f"scanned pages embed images below {PDF_SCAN_MIN_DPI:.0f} DPI "
                     f"(min {min_dpi:.0f} DPI). Low-resolution scans give poor "
                     "OCR; rescan at 200-300 DPI for documents."),
            details={"low_dpi_pages": [{"page": p, "dpi": round(d, 1)}
                                       for p, d in low_dpi_pages],
                     "pages_analyzed": len(blur_vars),
                     "threshold_dpi": PDF_SCAN_MIN_DPI,
                     "min_dpi": round(min_dpi, 1)},
        ))


def _check_pdf(path: Path, report: QualityReport, max_pages: int) -> None:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        report.add(QualityIssue(
            code="PDF_DEEP_CHECK_SKIPPED",
            severity=Severity.INFO,
            message=("PyMuPDF is not installed; skipping page count, "
                     "encryption, and text-extractability checks. "
                     "Install with: pip install pymupdf"),
        ))
        return

    try:
        doc = fitz.open(path)
    except Exception as exc:  # noqa: BLE001 - fitz raises many exception types
        report.add(QualityIssue(
            code="PDF_OPEN_FAILED",
            severity=Severity.ERROR,
            message=f"PDF could not be opened (likely corrupted): {exc}",
        ))
        return

    try:
        encrypted = bool(doc.needs_pass or doc.is_encrypted)
        if encrypted:
            report.add(QualityIssue(
                code="PDF_PASSWORD_PROTECTED",
                severity=Severity.ERROR,
                message=("PDF is password-protected. Content Understanding "
                         "rejects encrypted PDFs; remove the password first."),
            ))

        pages = doc.page_count
        report.metadata["page_count"] = pages
        if pages == 0:
            report.add(QualityIssue(
                code="PDF_NO_PAGES",
                severity=Severity.ERROR,
                message="PDF contains zero pages.",
            ))
        if pages > max_pages:
            report.add(QualityIssue(
                code="PDF_TOO_MANY_PAGES",
                severity=Severity.ERROR,
                message=(f"PDF has {pages} pages, exceeds the {max_pages}-page limit. "
                         "Split the document before processing."),
                details={"page_count": pages, "limit": max_pages},
            ))

        if encrypted or pages == 0:
            return

        if pages <= PDF_PAGE_SCAN_CAP:
            indices = list(range(pages))
        else:
            step = pages / PDF_PAGE_SCAN_CAP
            indices = sorted({int(i * step) for i in range(PDF_PAGE_SCAN_CAP)})

        per_page_chars: list[int] = []
        for i in indices:
            try:
                per_page_chars.append(len(doc.load_page(i).get_text("text")))
            except Exception:  # noqa: BLE001
                per_page_chars.append(0)

        sampled = len(per_page_chars)
        total_chars = sum(per_page_chars)
        avg_chars = total_chars / sampled if sampled else 0.0
        low_text_pages = sum(1 for c in per_page_chars
                             if c < LOW_TEXT_CHAR_THRESHOLD)
        empty_pages = sum(1 for c in per_page_chars if c < 5)
        low_ratio = low_text_pages / sampled if sampled else 0.0
        low_text_indices = [indices[i] for i, c in enumerate(per_page_chars)
                            if c < LOW_TEXT_CHAR_THRESHOLD]

        report.metadata["sampled_pages"] = sampled
        report.metadata["total_chars_sampled"] = total_chars
        report.metadata["avg_chars_per_page"] = round(avg_chars, 1)
        report.metadata["low_text_pages"] = low_text_pages
        report.metadata["empty_pages"] = empty_pages

        if avg_chars < 5:
            report.add(QualityIssue(
                code="PDF_LIKELY_SCANNED",
                severity=Severity.WARNING,
                message=("PDF appears to be image-only / scanned (almost no "
                         "extractable text on any sampled page). OCR will run, "
                         "which is slower and accuracy depends on scan quality."),
                details={"avg_chars_per_page": round(avg_chars, 1)},
            ))
        else:
            if avg_chars < MIN_AVG_CHARS_PER_PAGE:
                report.add(QualityIssue(
                    code="PDF_LOW_TEXT_DENSITY",
                    severity=Severity.WARNING,
                    message=(f"Average of {avg_chars:.0f} extractable characters "
                             f"per page is below the recommended minimum of "
                             f"{MIN_AVG_CHARS_PER_PAGE}. Pages may be mostly "
                             "images, blank, or weakly scanned."),
                    details={"avg_chars_per_page": round(avg_chars, 1),
                             "threshold": MIN_AVG_CHARS_PER_PAGE},
                ))
            if low_ratio >= LOW_TEXT_PAGE_RATIO and low_text_pages > 1:
                report.add(QualityIssue(
                    code="PDF_MANY_LOW_TEXT_PAGES",
                    severity=Severity.ERROR,
                    message=(f"{low_text_pages} of {sampled} sampled pages "
                             f"({low_ratio:.0%}) have fewer than "
                             f"{LOW_TEXT_CHAR_THRESHOLD} characters. Extraction "
                             "quality will be uneven across the document."),
                    details={"low_text_pages": low_text_pages,
                             "sampled_pages": sampled,
                             "char_threshold": LOW_TEXT_CHAR_THRESHOLD},
                ))
            if empty_pages > 0:
                report.add(QualityIssue(
                    code="PDF_BLANK_PAGES",
                    severity=Severity.INFO,
                    message=(f"{empty_pages} of {sampled} sampled pages look "
                             "completely blank (<5 characters)."),
                    details={"empty_pages": empty_pages,
                             "sampled_pages": sampled},
                ))

        _analyze_pdf_scan_quality(doc, low_text_indices, report)
    finally:
        doc.close()


def _check_image(path: Path, report: QualityReport) -> None:
    try:
        from PIL import Image, ImageFilter, ImageStat
    except ImportError:
        report.add(QualityIssue(
            code="IMAGE_DEEP_CHECK_SKIPPED",
            severity=Severity.INFO,
            message=("Pillow is not installed; skipping resolution and "
                     "corruption checks. Install with: pip install pillow"),
        ))
        return

    try:
        with Image.open(path) as img:
            img.verify()
        with Image.open(path) as img:
            width, height = img.size
            mode = img.mode
            gray = img.convert("L")
            laplacian = ImageFilter.Kernel(
                size=(3, 3),
                kernel=(0, 1, 0, 1, -4, 1, 0, 1, 0),
                scale=1,
                offset=0,
            )
            edges = gray.filter(laplacian)
            blur_variance = ImageStat.Stat(edges).var[0]
            stat = ImageStat.Stat(gray)
            mean_intensity = stat.mean[0]
            stdev_intensity = stat.stddev[0]
    except (OSError, ValueError) as exc:
        report.add(QualityIssue(
            code="IMAGE_OPEN_FAILED",
            severity=Severity.ERROR,
            message=f"Image could not be decoded (likely corrupted): {exc}",
        ))
        return
    except Exception as exc:  # noqa: BLE001
        report.add(QualityIssue(
            code="IMAGE_OPEN_FAILED",
            severity=Severity.ERROR,
            message=f"Image could not be decoded: {exc}",
        ))
        return

    report.metadata["width"] = width
    report.metadata["height"] = height
    report.metadata["mode"] = mode
    report.metadata["mean_intensity"] = round(mean_intensity, 1)
    report.metadata["stdev_intensity"] = round(stdev_intensity, 1)
    report.metadata["blur_variance"] = round(blur_variance, 1)

    if width < IMAGE_MIN_DIM or height < IMAGE_MIN_DIM:
        report.add(QualityIssue(
            code="IMAGE_TOO_SMALL",
            severity=Severity.ERROR,
            message=(f"Image is {width}x{height}px; minimum is "
                     f"{IMAGE_MIN_DIM}x{IMAGE_MIN_DIM}px."),
            details={"width": width, "height": height},
        ))
    if width > IMAGE_MAX_DIM or height > IMAGE_MAX_DIM:
        report.add(QualityIssue(
            code="IMAGE_TOO_LARGE",
            severity=Severity.ERROR,
            message=(f"Image is {width}x{height}px; maximum is "
                     f"{IMAGE_MAX_DIM}x{IMAGE_MAX_DIM}px. Resize before submitting."),
            details={"width": width, "height": height},
        ))

    if max(width, height) < 1000 and report.detected_kind == "image":
        report.add(QualityIssue(
            code="IMAGE_LOW_RESOLUTION",
            severity=Severity.WARNING,
            message=("Image resolution is low for OCR (long edge < 1000 px). "
                     "Consider rescanning at >= 150 DPI for documents."),
        ))

    if blur_variance < IMAGE_BLUR_VARIANCE_MIN:
        report.add(QualityIssue(
            code="IMAGE_BLURRY",
            severity=Severity.WARNING,
            message=(f"Image looks blurry or out of focus (variance of "
                     f"Laplacian = {blur_variance:.1f}, threshold "
                     f"{IMAGE_BLUR_VARIANCE_MIN}). OCR accuracy will suffer."),
            details={"blur_variance": round(blur_variance, 1),
                     "threshold": IMAGE_BLUR_VARIANCE_MIN},
        ))

    if stdev_intensity < IMAGE_LOW_CONTRAST_STDEV:
        if mean_intensity > 230:
            why = "image is almost entirely white (blank page?)"
        elif mean_intensity < 25:
            why = "image is almost entirely black"
        else:
            why = "pixel intensities have very little variation"
        report.add(QualityIssue(
            code="IMAGE_LOW_CONTRAST",
            severity=Severity.WARNING,
            message=(f"Low contrast: {why} (stdev={stdev_intensity:.1f}). "
                     "Text extraction will likely fail."),
            details={"mean_intensity": round(mean_intensity, 1),
                     "stdev_intensity": round(stdev_intensity, 1)},
        ))


def _check_office(path: Path, ext: str, report: QualityReport) -> None:
    if ext == ".docx":
        _check_docx(path, report)
    elif ext == ".xlsx":
        _check_xlsx(path, report)
    elif ext == ".pptx":
        _check_pptx(path, report)


def _check_docx(path: Path, report: QualityReport) -> None:
    try:
        import docx  # python-docx
    except ImportError:
        report.add(QualityIssue(
            code="DOCX_DEEP_CHECK_SKIPPED",
            severity=Severity.INFO,
            message=("python-docx is not installed; skipping character count. "
                     "Install with: pip install python-docx"),
        ))
        return

    try:
        document = docx.Document(str(path))
    except Exception as exc:  # noqa: BLE001
        msg = str(exc).lower()
        if "encrypted" in msg or "password" in msg:
            report.add(QualityIssue(
                code="DOCX_PASSWORD_PROTECTED",
                severity=Severity.ERROR,
                message="DOCX is password-protected. Remove the password first.",
            ))
        else:
            report.add(QualityIssue(
                code="DOCX_OPEN_FAILED",
                severity=Severity.ERROR,
                message=f"DOCX could not be opened (likely corrupted): {exc}",
            ))
        return

    total_chars = sum(len(p.text) for p in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                total_chars += len(cell.text)
    report.metadata["character_count"] = total_chars
    report.metadata["page_equivalent"] = -(-total_chars // CHARS_PER_TEXT_PAGE)

    if total_chars > OFFICE_MAX_CHARS:
        report.add(QualityIssue(
            code="OFFICE_TOO_MANY_CHARS",
            severity=Severity.ERROR,
            message=(f"DOCX contains {total_chars:,} characters; the limit is "
                     f"{OFFICE_MAX_CHARS:,}. Split the document."),
            details={"character_count": total_chars, "limit": OFFICE_MAX_CHARS},
        ))


def _check_xlsx(path: Path, report: QualityReport) -> None:
    try:
        import openpyxl
    except ImportError:
        report.add(QualityIssue(
            code="XLSX_DEEP_CHECK_SKIPPED",
            severity=Severity.INFO,
            message=("openpyxl is not installed; skipping sheet count. "
                     "Install with: pip install openpyxl"),
        ))
        return

    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=False)
    except Exception as exc:  # noqa: BLE001
        msg = str(exc).lower()
        if "encrypted" in msg or "password" in msg:
            report.add(QualityIssue(
                code="XLSX_PASSWORD_PROTECTED",
                severity=Severity.ERROR,
                message="XLSX is password-protected. Remove the password first.",
            ))
        else:
            report.add(QualityIssue(
                code="XLSX_OPEN_FAILED",
                severity=Severity.ERROR,
                message=f"XLSX could not be opened (likely corrupted): {exc}",
            ))
        return

    try:
        sheets = wb.sheetnames
        report.metadata["sheet_count"] = len(sheets)
        report.metadata["page_equivalent"] = len(sheets)
    finally:
        wb.close()


def _check_pptx(path: Path, report: QualityReport) -> None:
    try:
        from pptx import Presentation
    except ImportError:
        report.add(QualityIssue(
            code="PPTX_DEEP_CHECK_SKIPPED",
            severity=Severity.INFO,
            message=("python-pptx is not installed; skipping slide count. "
                     "Install with: pip install python-pptx"),
        ))
        return

    try:
        pres = Presentation(str(path))
    except Exception as exc:  # noqa: BLE001
        report.add(QualityIssue(
            code="PPTX_OPEN_FAILED",
            severity=Severity.ERROR,
            message=f"PPTX could not be opened (likely corrupted or encrypted): {exc}",
        ))
        return

    slide_count = len(pres.slides)
    report.metadata["slide_count"] = slide_count
    report.metadata["page_equivalent"] = slide_count


def _check_text(path: Path, report: QualityReport) -> None:
    try:
        raw = path.read_bytes()
    except OSError as exc:
        report.add(QualityIssue(
            code="READ_FAILED",
            severity=Severity.ERROR,
            message=f"Could not read file: {exc}",
        ))
        return

    decoded: Optional[str] = None
    for enc in ("utf-8", "utf-8-sig", "utf-16", "cp1252", "latin-1"):
        try:
            decoded = raw.decode(enc)
            report.metadata["encoding"] = enc
            break
        except UnicodeDecodeError:
            continue

    if decoded is None:
        report.add(QualityIssue(
            code="TEXT_UNDECODABLE",
            severity=Severity.WARNING,
            message=("Could not decode file with common encodings. "
                     "Re-save it as UTF-8 to avoid extraction errors."),
        ))
        return

    char_count = len(decoded)
    report.metadata["character_count"] = char_count
    report.metadata["page_equivalent"] = max(1, -(-char_count // CHARS_PER_TEXT_PAGE))

    if char_count > TEXT_MAX_CHARS:
        report.add(QualityIssue(
            code="TEXT_TOO_MANY_CHARS",
            severity=Severity.ERROR,
            message=(f"File contains {char_count:,} characters; the limit is "
                     f"{TEXT_MAX_CHARS:,}."),
        ))
    if char_count == 0:
        report.add(QualityIssue(
            code="TEXT_EMPTY",
            severity=Severity.ERROR,
            message="File is empty after decoding.",
        ))


# ---------------------------------------------------------------------------
# Top-level orchestrator
# ---------------------------------------------------------------------------

def check_document(path: str | os.PathLike[str], *, mode: str = "standard") -> QualityReport:
    """Run all applicable quality checks against ``path``.

    Parameters
    ----------
    path : str | PathLike
        File on disk to inspect.
    mode : "standard" | "pro"
        Apply the standard limits (default) or the tighter Pro-mode preview
        limits (100 MB / 150 pages, PDF + image-only).
    """
    if mode not in {"standard", "pro"}:
        raise ValueError(f"mode must be 'standard' or 'pro', got {mode!r}")

    p = Path(path)
    report = QualityReport(
        file_path=str(p),
        file_size_bytes=0,
        extension="",
        detected_kind="unknown",
        mode=mode,
        passed=True,
    )

    if not p.exists():
        report.add(QualityIssue(
            code="FILE_NOT_FOUND",
            severity=Severity.ERROR,
            message=f"File does not exist: {p}",
        ))
        return report
    if not p.is_file():
        report.add(QualityIssue(
            code="NOT_A_FILE",
            severity=Severity.ERROR,
            message=f"Path is not a regular file: {p}",
        ))
        return report

    try:
        size = p.stat().st_size
    except OSError as exc:
        report.add(QualityIssue(
            code="STAT_FAILED",
            severity=Severity.ERROR,
            message=f"Could not stat file: {exc}",
        ))
        return report

    report.file_size_bytes = size
    if size == 0:
        report.add(QualityIssue(
            code="FILE_EMPTY",
            severity=Severity.ERROR,
            message="File is zero bytes.",
        ))
        return report

    ext = p.suffix.lower()
    report.extension = ext
    kind = _detect_kind(ext)
    report.detected_kind = kind

    mime, _ = mimetypes.guess_type(p.name)
    if mime:
        report.metadata["mime_type"] = mime

    if kind == "unknown":
        report.add(QualityIssue(
            code="UNSUPPORTED_EXTENSION",
            severity=Severity.ERROR,
            message=(f"Extension {ext!r} is not in the Content Understanding "
                     f"supported list ({', '.join(sorted(ALL_SUPPORTED_EXT))})."),
        ))
        return report

    if mode == "pro" and ext not in PRO_SUPPORTED_EXT:
        report.add(QualityIssue(
            code="PRO_MODE_UNSUPPORTED_EXTENSION",
            severity=Severity.ERROR,
            message=(f"Pro mode only accepts PDF/TIFF/image. {ext!r} is not allowed. "
                     "Switch to standard mode or convert the file."),
        ))

    size_limit = {
        "pdf":   PRO_MAX_BYTES if mode == "pro" else DOC_IMAGE_MAX_BYTES,
        "image": PRO_MAX_BYTES if mode == "pro" else DOC_IMAGE_MAX_BYTES,
        "office": OFFICE_MAX_BYTES,
        "text":  TEXT_MAX_BYTES,
    }[kind]
    if size > size_limit:
        report.add(QualityIssue(
            code="FILE_TOO_LARGE",
            severity=Severity.ERROR,
            message=(f"File is {size / MB:.2f} MB; the limit for "
                     f"{kind} files in {mode} mode is {size_limit / MB:.0f} MB."),
            details={"size_bytes": size, "limit_bytes": size_limit},
        ))

    _check_magic(p, ext, report)

    max_pages = PRO_MAX_PAGES if mode == "pro" else DOC_IMAGE_MAX_PAGES
    if kind == "pdf":
        _check_pdf(p, report, max_pages=max_pages)
    elif kind == "image":
        _check_image(p, report)
    elif kind == "office":
        _check_office(p, ext, report)
    elif kind == "text":
        _check_text(p, report)

    return report
